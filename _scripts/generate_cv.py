#!/usr/bin/env python3
"""Génère un CV archivé pour une candidature, à partir d'un CV de base.

Usage:
    python generate_cv.py --base ../_base/CV_base.docx --poste "AI Engineer" \
        --entreprise UNICEF --categorie AI-ML-Engineering --annee 2026 \
        --lien "https://..." --statut "Envoyée"
"""

import argparse
import json
import re
import shutil
import sys
import unicodedata
from datetime import date
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parent.parent
CONFIG_PATH = Path(__file__).resolve().parent / "config.json"
LOG_PATH = ROOT / "CANDIDATURES.md"
DEFAULT_BASE = ROOT / "_base" / "CV_base.docx"

CATEGORIES = [
    "AI-ML-Engineering",
    "Autre",
    "Charge-Etudes-Marketing",
    "Data-Science",
    "LLM-Engineering",
    "MEAL-Suivi-Evaluation",
]


def sanitize(text: str) -> str:
    """Retire accents/espaces/caractères spéciaux, joint les mots par des tirets."""
    normalized = unicodedata.normalize("NFKD", text)
    ascii_text = normalized.encode("ascii", "ignore").decode("ascii")
    ascii_text = re.sub(r"[^A-Za-z0-9\s-]", "", ascii_text)
    words = ascii_text.split()
    return "-".join(words)


def load_config() -> dict:
    if not CONFIG_PATH.exists():
        sys.exit(f"Config introuvable: {CONFIG_PATH}")
    config = json.loads(CONFIG_PATH.read_text(encoding="utf-8"))
    if config.get("prenom") == "Prenom" or config.get("nom") == "Nom":
        sys.exit(
            f"Renseigne ton prénom/nom dans {CONFIG_PATH} avant de générer un CV."
        )
    return config


def append_log(row: list[str]) -> None:
    if not LOG_PATH.exists():
        sys.exit(f"Log introuvable: {LOG_PATH}")
    with LOG_PATH.open("a", encoding="utf-8", newline="\n") as f:
        f.write("| " + " | ".join(row) + " |\n")


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument(
        "--base",
        default=str(DEFAULT_BASE),
        help="Chemin du CV de base (.docx). Par défaut: _base/CV_base.docx",
    )
    parser.add_argument("--poste", required=True, help="Intitulé du poste visé")
    parser.add_argument("--entreprise", required=True, help="Nom de l'entreprise")
    parser.add_argument("--categorie", required=True, choices=CATEGORIES)
    parser.add_argument(
        "--annee", type=int, default=date.today().year, help="Année de candidature"
    )
    parser.add_argument("--lien", default="", help="Lien vers l'offre")
    parser.add_argument("--statut", default="Envoyée", help="Statut de la candidature")
    parser.add_argument(
        "--titre",
        default="",
        help="Remplace l'accroche du CV (2e ligne) par ce titre, adapté à l'offre",
    )
    args = parser.parse_args()

    base_path = Path(args.base)
    if not base_path.exists():
        sys.exit(f"CV de base introuvable: {base_path}")
    if base_path.suffix.lower() != ".docx":
        sys.exit("Le CV de base doit être un fichier .docx")

    config = load_config()
    prenom = sanitize(config["prenom"])
    nom = sanitize(config["nom"])
    poste = sanitize(args.poste)
    entreprise = sanitize(args.entreprise)

    filename = f"{prenom}_{nom}_{poste}_{entreprise}_{args.annee}.docx"
    dest_dir = ROOT / args.categorie
    dest_dir.mkdir(exist_ok=True)
    dest_path = dest_dir / filename

    if dest_path.exists():
        sys.exit(f"Un CV existe déjà pour cette candidature: {dest_path}")

    shutil.copy2(base_path, dest_path)

    if args.titre:
        doc = Document(dest_path)
        tagline_paragraph = doc.paragraphs[1]
        for run in tagline_paragraph.runs:
            run.text = ""
        if tagline_paragraph.runs:
            tagline_paragraph.runs[0].text = args.titre
        else:
            tagline_paragraph.add_run(args.titre)
        doc.save(dest_path)

    rel_path = dest_path.relative_to(ROOT).as_posix()
    # Ordre des colonnes = celui déjà en place dans CANDIDATURES.md
    append_log(
        [
            date.today().isoformat(),
            args.poste,
            args.entreprise,
            args.categorie,
            args.lien,
            rel_path,
            args.statut,
        ]
    )

    print(f"CV cree: {dest_path}")
    print(f"Logge dans: {LOG_PATH}")


if __name__ == "__main__":
    main()
