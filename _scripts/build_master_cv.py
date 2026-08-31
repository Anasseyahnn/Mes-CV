from docx import Document
from docx.shared import Pt, RGBColor
import os

path = os.path.join(os.path.dirname(__file__), "..", "_base", "CV_base.docx")

d = Document()

style = d.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10.5)


def add_name(text):
    p = d.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = RGBColor(0x1F, 0x2D, 0x50)


def add_tagline(text):
    p = d.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(12.5)
    r.font.color.rgb = RGBColor(0x2E, 0x5E, 0x8C)


def add_contact(text):
    p = d.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor(0x40, 0x40, 0x40)


def add_section(text):
    p = d.add_paragraph()
    r = p.add_run(text.upper())
    r.bold = True
    r.font.size = Pt(11.5)
    r.font.color.rgb = RGBColor(0x1F, 0x2D, 0x50)
    pf = p.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after = Pt(3)


def add_body(text):
    p = d.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(6)


def add_job(title_company, dates):
    p = d.add_paragraph()
    r = p.add_run(title_company)
    r.bold = True
    r.font.size = Pt(10.5)
    if dates:
        tab = p.add_run("\t" + dates)
        tab.italic = True
        tab.font.size = Pt(9.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)


def add_bullet(text):
    p = d.add_paragraph(style='List Bullet')
    r = p.add_run(text)
    r.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(2)


add_name("Anassé Yahanan")
add_tagline("Lead Data Scientist & AI Engineer — Stratégie Data, Machine Learning & IA Générative")
add_contact(
    "Abidjan, Côte d'Ivoire  |  anasseyahanan@gmail.com  |  "
    "https://anasseyahnn.github.io/portfolio/  |  https://github.com/Anasseyahnn  |  "
    "https://linkedin.com/in/anasse-yahanan-bouagba-3b39aa242/"
)
add_body("3+ ans d'expérience  ·  15+ projets IA livrés  ·  50+ professionnels formés  ·  Fondateur d'un SaaS IA")

add_section("Résumé")
add_body(
    "Data Scientist / AI Engineer avec 3+ ans d'expérience dans le pilotage de projets data de bout en bout — "
    "de la stratégie à la mise en production — pour des études de marché et des programmes de développement en "
    "Afrique de l'Ouest. Fondateur de BRVM AI, plateforme SaaS d'analyse boursière propulsée par l'IA, et concepteur "
    "d'un programme de formation certifiant \"Ingénieur LLM\" (10 modules, 250h, 50+ professionnels formés). "
    "Expertise combinée en machine learning (Scikit-Learn, PyTorch), IA générative (LLM, RAG, LangChain, LlamaIndex) "
    "et industrialisation (FastAPI, Docker, Kubernetes). Actuellement Chargé d'Études Statistique & Marketing Senior "
    "chez OpinionWay Afrique Subsaharienne, où je pilote 10+ études de marché de bout en bout pour des clients "
    "grands comptes. Recherche un poste à responsabilité (Lead / Head / Manager) en Data Science, IA ou Ingénierie LLM."
)

add_section("Compétences clés")
add_bullet("Leadership & pilotage : pilotage de projets data de bout en bout · gestion de parties prenantes grands comptes · formation et montée en compétence d'équipes · fondation et direction produit (SaaS)")
add_bullet("Machine learning & IA : Scikit-Learn · PyTorch · LLM · RAG · LangChain · LlamaIndex · Ollama · Groq")
add_bullet("Ingénierie & déploiement : Python · R · SQL · FastAPI · Docker · Kubernetes")
add_bullet("Visualisation & data storytelling : Power BI · ggplot2 · Quarto · Streamlit · Next.js")

add_section("Expérience professionnelle")

add_job("Chargé d'Études Statistique & Marketing Senior — OpinionWay Afrique Subsaharienne", "2025 – Présent")
add_bullet("Pilote 10+ études de marché de bout en bout pour des clients grands comptes en Afrique de l'Ouest")
add_bullet("Optimise la segmentation clientèle via clustering K-means (−40% de temps d'analyse)")
add_bullet("Automatise les pipelines ETL en R & Python (−40% de temps de reporting)")
add_bullet("Conçoit des dashboards Power BI adoptés par les décideurs pour le suivi des indicateurs clés")

add_job("Consultant Chargé d'Étude Statistique & Marketing — Cabinet InsightPlus", "2024")
add_bullet("Réalise 3 études de marché complètes en 3 mois (régressions, tests d'hypothèses), en autonomie totale")
add_bullet("Formule des recommandations stratégiques fondées sur les données pour une clientèle diversifiée")

add_job("Analyste Revenu Assurance Junior — Banque Atlantique Côte d'Ivoire", "2022")
add_bullet("Contrôle des données financières mensuelles, fiabilité 100% des rapports produits pour la direction")
add_bullet("Détecte et corrige des écarts comptables récurrents en environnement bancaire réglementé")

add_section("Leadership, produit & formation")
add_job("Fondateur — BRVM AI, SaaS d'analyse boursière propulsé par l'IA", "")
add_bullet("Conçoit, développe et déploie en autonomie complète une plateforme SaaS d'analyse financière de la Bourse Régionale (Next.js, Vercel), LLM intégrés pour des insights à latence <150ms")
add_job("Concepteur & Formateur — Programme certifiant \"Ingénieur LLM\"", "")
add_bullet("Conçoit et anime un programme de 10 modules / 250h (2 parcours dev/analyste, capstone déployé, évaluation automatique par LLM-as-judge) — 50+ professionnels formés au LLM, RAG et déploiement d'agents IA")
add_job("Consultant Formateur R & Python — INSSEDS", "")
add_bullet("Conçoit et anime des formations professionnelles en R et Python pour l'analyse et la visualisation de données")
add_job("Suivi-Évaluation (MEAL) — Projets financés USAID et Ambassade des États-Unis", "")
add_bullet("Collecte, contrôle qualité et analyse de données de suivi de programmes de développement, restitution à des bailleurs internationaux")

add_section("Projets notables (portfolio GitHub)")
add_bullet("BRVM AI (SaaS) — Analyse financière et boursière automatisée par l'IA (Next.js, Vercel, LLM, RAG)")
add_bullet("Labo Ingénieur LLM — Plateforme de formation LLM : 10 modules, rendus d'exercices, évaluation automatique LLM-as-judge (FastAPI, Next.js)")
add_bullet("AWASQA — SaaS de transcription, codage thématique IA et reporting automatisé d'entretiens qualitatifs (TypeScript, multi-LLM)")
add_bullet("Sondaq — Plateforme no-code d'analyse d'études de marché : segmentation, positionnement (Next.js, TypeScript)")
add_bullet("SQL-Chat — Traducteur langage naturel → SQL avec auto-charting et analyses narratives (Python, Ollama)")
add_bullet("Assistant_droit_ivoirien_AI — Assistant juridique IA (RAG) spécialisé Code du Travail ivoirien")
add_bullet("Vision OCR Pro — Extraction et structuration de texte depuis des images, 100% local via Ollama")
add_bullet("Groq R Integration — Package connectant des scripts statistiques R aux puces LPU Groq")
add_bullet("cancer_du_sein_predict — Diagnostic prédictif du cancer du sein par machine learning (Streamlit)")

add_section("Formation")
add_body("Formation continue et autoformation en Data Science, Ingénierie de données et Intelligence Artificielle appliquée. Détails et certifications disponibles sur le portfolio et le profil LinkedIn.")

d.save(path)
print("saved:", path)
